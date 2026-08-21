import os
import io
import uuid
import logging
from datetime import datetime
from pathlib import Path
from typing import List, Tuple
from sqlalchemy.orm import Session
from sqlalchemy.exc import SQLAlchemyError
import openpyxl
from openpyxl.styles import Font, PatternFill, Alignment

from app.models.question import Question
from app.models.system import Export
from app.core.storage import get_export_dir

logger = logging.getLogger(__name__)

def _sanitize_for_excel(text: str) -> str:
    """Ngăn chặn Excel Formula Injection"""
    if not text:
        return ""
    text = str(text)
    if text.startswith(("=", "+", "-", "@")):
        return f"'{text}"
    return text

def _generate_excel_workbook(questions: List[Question]) -> openpyxl.Workbook:
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Câu hỏi"

    # Tính số cột option tối đa (ít nhất là 4 để giữ form A, B, C, D)
    max_opts = 4
    if questions:
        max_opts = max(max([len(q.options) for q in questions] + [4]), 4)

    # 1. Tạo Header
    headers = ["STT", "Nội dung câu hỏi"]
    
    # Cột đáp án 
    for i in range(max_opts):
        letter = chr(65 + i) # A, B, C...
        headers.append(f"Đáp án {letter}")
        
    headers.extend([
        "Đáp án đúng", 
        "Giải thích", 
        "Độ khó", 
        "Bloom", 
        "Loại câu hỏi", 
        "Nguồn"
    ])

    ws.append(headers)

    # Format Header
    header_font = Font(color="FFFFFF", bold=True)
    header_fill = PatternFill(start_color="000080", end_color="000080", fill_type="solid")
    header_alignment = Alignment(horizontal="center", vertical="center")

    for col_num in range(1, len(headers) + 1):
        cell = ws.cell(row=1, column=col_num)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_alignment

    # Bật AutoFilter và Freeze
    ws.auto_filter.ref = f"A1:{openpyxl.utils.get_column_letter(len(headers))}1"
    ws.freeze_panes = "A2"

    # 2. Điền Data
    data_alignment = Alignment(vertical="top", wrap_text=True)

    for idx, q in enumerate(questions, 1):
        row_data = [
            idx,
            _sanitize_for_excel(q.content)
        ]

        sorted_opts = sorted(q.options, key=lambda x: x.id)
        
        correct_letters = []
        # Điền nội dung options
        for i in range(max_opts):
            if i < len(sorted_opts):
                opt = sorted_opts[i]
                row_data.append(_sanitize_for_excel(opt.content))
                if opt.is_correct:
                    correct_letters.append(chr(65 + i))
            else:
                row_data.append("") # padding

        # Đáp án đúng
        row_data.append(", ".join(correct_letters))
        
        # Các cột còn lại
        row_data.extend([
            _sanitize_for_excel(q.explanation),
            q.difficulty or "",
            q.bloom_level or "",
            q.question_type or "",
            _sanitize_for_excel(q.material.title if q.material else "N/A")
        ])

        ws.append(row_data)

        for col_num in range(1, len(headers) + 1):
            ws.cell(row=idx + 1, column=col_num).alignment = data_alignment

    # 3. Chỉnh Auto-width 
    for col in ws.columns:
        max_length = 0
        column = col[0].column_letter 
        for cell in col:
            try: 
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        
        adjusted_width = min((max_length + 2), 50)
        ws.column_dimensions[column].width = adjusted_width

    return wb

def export_questions_to_excel(db: Session, questions: List[Question], user_id: int, question_ids: List[int]) -> Tuple[Path, str]:
    course_ids = {q.course_id for q in questions if q.course_id is not None}
    final_course_id = course_ids.pop() if len(course_ids) == 1 else None

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    unique_id = uuid.uuid4().hex
    filename = f"questions_export_{timestamp}_{unique_id}.xlsx"
    
    export_dir = get_export_dir()
    os.makedirs(export_dir, exist_ok=True)
    
    final_path = export_dir / filename
    temp_path = export_dir / f"{filename}.tmp"
    
    wb = _generate_excel_workbook(questions)
    
    try:
        wb.save(temp_path)
        os.replace(temp_path, final_path)
    except Exception as e:
        if os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except Exception:
                pass
        raise e
        
    try:
        new_export = Export(
            course_id=final_course_id,
            exported_by=user_id,
            file_path=filename,
            format="xlsx",
            question_ids=question_ids
        )
        db.add(new_export)
        db.flush()
        db.commit()
    except Exception as e:
        db.rollback()
        try:
            os.remove(final_path)
        except Exception as cleanup_err:
            logger.error(f"Failed to cleanup file {final_path} after DB error: {cleanup_err}")
        raise e
        
    return final_path, filename

def export_questions_to_word(db: Session, questions: List[Question], user_id: int, question_ids: List[int]) -> Tuple[Path, str]:
    import docx
    
    course_ids = {q.course_id for q in questions if q.course_id is not None}
    final_course_id = course_ids.pop() if len(course_ids) == 1 else None

    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    unique_id = uuid.uuid4().hex
    filename = f"questions_export_{timestamp}_{unique_id}.docx"
    
    export_dir = get_export_dir()
    os.makedirs(export_dir, exist_ok=True)
    
    final_path = export_dir / filename
    temp_path = export_dir / f"{filename}.tmp"
    
    doc = docx.Document()
    doc.add_heading("Danh sách câu hỏi", level=1)
    
    for idx, q in enumerate(questions, 1):
        doc.add_heading(f"Câu {idx}: {q.content}", level=2)
        
        sorted_opts = sorted(q.options, key=lambda x: x.id)
        for i, opt in enumerate(sorted_opts):
            letter = chr(65 + i)
            p = doc.add_paragraph()
            run = p.add_run(f"{letter}. {opt.content}")
            if opt.is_correct:
                run.bold = True
                run.underline = True
        
        p_meta = doc.add_paragraph()
        p_meta.add_run(f"Độ khó: {q.difficulty or 'N/A'} | Bloom: {q.bloom_level or 'N/A'}").italic = True
        
        if q.explanation:
            doc.add_paragraph(f"Giải thích: {q.explanation}")
            
        doc.add_paragraph() # Dòng trống
    
    try:
        doc.save(temp_path)
        os.replace(temp_path, final_path)
    except Exception as e:
        if os.path.exists(temp_path):
            try:
                os.remove(temp_path)
            except Exception:
                pass
        raise e
        
    try:
        new_export = Export(
            course_id=final_course_id,
            exported_by=user_id,
            file_path=filename,
            format="docx",
            question_ids=question_ids
        )
        db.add(new_export)
        db.flush()
        db.commit()
    except Exception as e:
        db.rollback()
        try:
            os.remove(final_path)
        except Exception as cleanup_err:
            logger.error(f"Failed to cleanup file {final_path} after DB error: {cleanup_err}")
        raise e
        
    return final_path, filename
